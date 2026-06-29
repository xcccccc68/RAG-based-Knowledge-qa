from typing import List, Dict
import docx
import base64
import io
import re
from utils.logger import logger


class DocxProcessor:
    """
    Word文档处理器，负责从DOCX文档中提取文本、表格和图片元素
    """
    
    def __init__(self):
        """
        初始化Word文档处理器
        """
        pass
    
    def extract_elements(self, docx_path: str) -> List[Dict]:
        """
        从DOCX中提取文本、表格和图片元素

        Args:
            docx_path: DOCX文件路径

        Returns:
            List[Dict]: 包含所有元素的列表，每个元素包含类型、内容、位置等信息
        """
        logger.info(f"开始处理DOCX文件: {docx_path}")
        elements = []
        doc = docx.Document(docx_path)

        # 提取文本和表格
        logger.info(f"提取文本段落")
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                elements.append({
                    'type': 'text',
                    'content': paragraph.text,
                    'position': f'paragraph_{para_idx}'
                })

        # 提取表格
        logger.info(f"提取表格，共{len(doc.tables)}个表格")
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

            # 提取表格标题
            table_title = self.extract_table_title(elements, len(elements))
            elements.append({
                'type': 'table',
                'table_data': table_data,
                'title': table_title,
                'position': f'table_{table_idx}'
            })

        # 提取图片
        logger.info(f"提取图片")
        image_count = 0
        for rel_idx, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                try:
                    # 读取图片
                    image_part = doc.part.rels[rel.rId].target_part
                    image_bytes = image_part.blob

                    # 编码为base64
                    img_base64 = base64.b64encode(image_bytes).decode('utf-8')

                    # 提取图片标题
                    image_title = self.extract_image_title(elements, len(elements))

                    elements.append({
                        'type': 'image',
                        'image_bytes': io.BytesIO(image_bytes),
                        'image_base64': img_base64,
                        'title': image_title,
                        'position': f'image_{rel_idx}'
                    })
                    image_count += 1
                except Exception as e:
                    logger.error(f"提取图片失败: {e}")
                    continue

        logger.info(f"DOCX文件处理完成，共提取{len(elements)}个元素，其中图片{image_count}个")
        return elements
    
    def extract_table_title(self, elements: List[Dict], table_idx: int) -> str:
        """
        提取表格标题
        
        Args:
            elements: 所有元素列表
            table_idx: 表格索引
            
        Returns:
            str: 表格标题
        """
        # 查找表格前的文本元素，寻找可能的标题
        for i in range(table_idx - 1, -1, -1):
            if elements[i]['type'] == 'text':
                text = elements[i]['content']
                if re.search(r'表\s*\d+.*', text) or re.search(r'表格\s*\d*.*', text):
                    return text
                # 如果找到非空文本且不是表格，也作为标题候选
                if text.strip():
                    return text
            elif elements[i]['type'] != 'text':
                break
        return ""
    
    def extract_image_title(self, elements: List[Dict], image_idx: int) -> str:
        """
        提取图片标题
        
        Args:
            elements: 所有元素列表
            image_idx: 图片索引
            
        Returns:
            str: 图片标题
        """
        # 查找图片前的文本元素，寻找可能的标题
        for i in range(image_idx - 1, -1, -1):
            if elements[i]['type'] == 'text':
                text = elements[i]['content']
                if re.search(r'图\s*\d+.*', text) or re.search(r'图片\s*\d*.*', text):
                    return text
                # 如果找到非空文本且不是图片，也作为标题候选
                if text.strip():
                    return text
            elif elements[i]['type'] != 'text':
                break
        return ""
